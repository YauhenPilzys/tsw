import os
import re
import zipfile
import ftplib
import io
import urllib.parse
from datetime import datetime, timedelta
from ftplib import FTP
from io import BytesIO
from django.shortcuts import render, get_object_or_404
from django.conf import settings
from django.template.loader import render_to_string
from django.contrib.auth import authenticate, get_user_model
from django.contrib.auth.hashers import make_password
from django.contrib.contenttypes.models import ContentType
from django.forms.models import model_to_dict
from django.db.models import Q
from django.utils.timezone import make_aware
from django.http import HttpResponse, FileResponse

from rest_framework import status, permissions, viewsets
from rest_framework.viewsets import ModelViewSet, ReadOnlyModelViewSet
from rest_framework.views import APIView
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.exceptions import NotFound
from rest_framework.permissions import IsAuthenticated, IsAdminUser, BasePermission
from rest_framework.filters import SearchFilter, OrderingFilter
from rest_framework_simplejwt.views import TokenObtainPairView
from django_filters.rest_framework import DjangoFilterBackend

from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from weasyprint import HTML, CSS

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import *
from .serializers import *
from .paginations import AllOtherAPIListPagination
from .directory import *
from .mixins import *
from .xml_uved import generate_xml


import requests
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
import xml.etree.ElementTree as ET
from django.db import transaction
from django.utils.dateparse import parse_date
from django.db.models import Q
from django.utils.timezone import make_aware, utc
User = get_user_model()

#http://178.168.146.114:8770/reset-password/
# {
#     "username": "admin",
#     "old_password": "root",
#     "new_password": "admin"
# }
class ResetPasswordView(APIView):
    """Позволяет пользователю сменить пароль при входе, без токена"""

    authentication_classes = []
    permission_classes = []

    def post(self, request):
        """Меняем пароль по `username` и `old_password`"""
        username = request.data.get("username")
        old_password = request.data.get("old_password")
        new_password = request.data.get("new_password")

        if not username or not old_password or not new_password:
            return Response({"error": "Все поля обязательны"}, status=status.HTTP_400_BAD_REQUEST)

        # Проверяем, существует ли пользователь
        try:
            user = User.objects.get(username=username)
        except User.DoesNotExist:
            return Response({"error": "Пользователь не найден"}, status=status.HTTP_404_NOT_FOUND)

        # Аутентифицируем пользователя по старому паролю
        user_check = authenticate(username=username, password=old_password)
        if user_check is None:
            return Response({"error": "Неверный старый пароль"}, status=status.HTTP_400_BAD_REQUEST)

        # Обновляем пароль
        user.password = make_password(new_password)
        user.save()

        return Response({
            "success": "Пароль успешно изменён",
            "username": user.username,
            "new_password": new_password
        }, status=status.HTTP_200_OK)


class SendXMLFile(APIView):  # Отправка XML на FTP
    permission_classes = (IsAuthenticated,)

    def get(self, request, *args, **kwargs):
        # Получаем объект Notice
        notice = get_object_or_404(Notice, id=kwargs['pk'])

        # Генерация XML
        try:
            xml_data = generate_xml(notice)
        except Exception as e:
            return Response({'ERROR': f'Ошибка генерации XML: {e}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # Формируем имя файла
        guid_upper = notice.guid.upper()
        file_name = f"B_{guid_upper}-{notice.order.warehouse.customs_post}_UV.xml"

        # Отправка файла на FTP
        try:
            with ftplib.FTP() as ftp:
                ftp.connect("178.168.146.109", 21, timeout=60)
                ftp.set_pasv(False)
                ftp.login(user="alestaftp", passwd="123456_qaz")
                ftp.cwd('/')
                ftp.storbinary(f"STOR {file_name}", io.BytesIO(xml_data))
                ftp.quit()
        except ftplib.all_errors as e:
            return Response({'ERROR': f'Ошибка при отправке файла на FTP: {e}'},
                            status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # Статус заказа НЕ меняем после отправки (остается "в работе")

        # Запись в EClient
        transport = notice.transport_notice.first()
        regN_DOK = Doc.objects.filter(notice=notice, doc_code="02015").values_list("doc_number", flat=True).first()

        try:
            EClient.objects.using('second_db').create(
                regN_TS=transport.number.upper() if transport else "",
                regN_DOK=regN_DOK or "",
                countKod=0,
                guidXML=notice.guid.upper(),
                id_us=68,
                tip=4,
                status="в очереди на отправку",
                statID=-1,
                postIN=notice.order.warehouse.customs_post if notice.order.warehouse else "16466",
                id_xmlf=0,
                regID=None,
                lc=localtime(now())
            )
        except Exception as e:
            return Response({'ERROR': f'Ошибка при создании записи в EClient: {e}'},
                            status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        return Response({'OK': f'Файл {file_name} успешно отправлен.'}, status=status.HTTP_200_OK)



class CheckEClientStatus(APIView):  # Проверка статуса в EClient
    permission_classes = (IsAuthenticated,)

    def get(self, request, *args, **kwargs):
        notice = get_object_or_404(Notice, id=kwargs['pk'])

        try:
            # Получаем запись из EClient по GUID XML
            eclient_record = EClient.objects.using('second_db').filter(guidXML=notice.guid.upper()).first()

            if eclient_record:
                eclient_status = eclient_record.status.lower()  # Приводим к нижнему регистру для проверки

                if "отказ в регистрации" in eclient_status:
                    # Если в статусе есть "отказ в регистрации", меняем заказ на "ошибка при отправке" (2)
                    notice.order.status_order = '2'
                    notice.order.save(update_fields=['status_order'])
                    return Response({'ERROR': 'Отказ в регистрации. Статус заказа обновлен на 2 (ошибка при отправке).'},
                                    status=status.HTTP_400_BAD_REQUEST)

                elif "зарегистрировано" in eclient_status or "снято с контроля" in eclient_status:
                    # Если уведомление зарегистрировано или снято с контроля, заказ = "отправлен" (3)
                    notice.order.status_order = '3'
                    notice.order.save(update_fields=['status_order'])
                    return Response({'OK': 'Уведомление зарегистрировано в НАСЭД. Статус заказа обновлен на 3 (отправлен).'},
                                    status=status.HTTP_200_OK)

                return Response({'INFO': f'Статус в EClient: {eclient_status}. Ожидание регистрации.'},
                                status=status.HTTP_202_ACCEPTED)

            return Response({'INFO': 'Ответ от НАСЭД еще не получен.'}, status=status.HTTP_202_ACCEPTED)

        except Exception as e:
            return Response({'ERROR': f'Ошибка при получении статуса из EClient: {e}'},
                            status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# Скачивание XML
# http://127.0.0.1:8000/notices/1/download_xml/
def download_notice_xml(request, notice_id):
    try:
        # Получаем объект Notice по id_notice
        notice = get_object_or_404(Notice, id=notice_id)

        # Генерируем XML
        try:
            xml_data = generate_xml(notice)
        except Exception as e:
            return JsonResponse({'error': f'Ошибка генерации XML: {str(e)}'}, status=500)

        # Формируем имя файла
        guid_upper = notice.guid.upper()  # GUID в верхнем регистре
        file_name = f"B_{guid_upper}-{notice.order.warehouse.customs_post}_UV.xml"

        # Возвращаем файл для скачивания
        response = HttpResponse(xml_data, content_type="application/xml")
        response['Content-Disposition'] = f'attachment; filename="{file_name}"'
        return response

    except Notice.DoesNotExist:
        # Обработка, если уведомление не найдено
        return JsonResponse({'error': 'Уведомление с указанным ID не существует'}, status=404)

    except Exception as e:
        # Общая обработка ошибок
        return JsonResponse({'error': f'Произошла ошибка: {str(e)}'}, status=500)




# Запрос на создание пропуска на заезд
# http://127.0.0.1:8000/pass/1/ - 1 - id logbook
def generate_pass_pdf(request, logbook_id):
    try:
        # Получение записи LogBook по ID
        logbook = LogBook.objects.select_related('place_park', 'warehouse').get(pk=logbook_id)
    except LogBook.DoesNotExist:
        return HttpResponse("Запись в журнале не найдена", status=404)

    # Абсолютный URL логотипа
    logo_url = request.build_absolute_uri(f"{settings.MEDIA_URL}images/logo.png")

    # Формирование строки номера ТС с учетом прицепа
    if logbook.trailer_number:
        vehicle_info = f"{logbook.vehicle_number}/{logbook.trailer_number}"
        filename = f"{logbook.vehicle_number}_{logbook.trailer_number}_{logbook.id:07d}.pdf"
    else:
        vehicle_info = f"{logbook.vehicle_number}"
        filename = f"{logbook.vehicle_number}_{logbook.id:07d}.pdf"

    # Подготовка данных для шаблона
    context = {
        "logbook": logbook,
        "vehicle_info": vehicle_info,  # Добавляем в контекст
        "rules": [
            "проносить на территорию товары, которые могут нанести ущерб здоровью либо товарам, хранящимся на СВХ (оружие, боеприпасы, взрывчатые, отравляющие вещества, горючие и легковоспламеняющиеся материалы, спиртные напитки и т.д.);",
            "курить, справлять естественные надобности вне установленных мест;",
            "пользоваться открытым огнем с целью разогрева пищи и в иных случаях;",
            "приносить, распивать спиртные напитки, употреблять психотропные и наркотические вещества;",
            "выбрасывать мусор, пищевые отходы вне специально отведённых для этих целей мест;",
            "производить техническое обслуживание, ремонт и мойку транспортных средств, иные виды работ, которые могут привести к порче асфальтированного покрытия и нанесению физического вреда водителю (иным лицам);",
            "допускать вытекание горюче-смазочных материалов из автомобиля;",
            "управлять транспортным средством в ЗТК в состоянии алкогольного, наркотического либо иного токсического опьянения и (или) передавать управление транспортного средства лицу, находящемуся в аналогичном состоянии;",
            "самостоятельно, без предварительного уведомления специалиста СВХ и разрешения таможенного органа покидать территорию ЗТК и (или) выезжать на транспортном средстве."
        ],
        "logo_path": logo_url
    }

    # Генерация HTML из шаблона
    html_string = render_to_string('pass_template.html', context)

    # Создание PDF
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'inline; filename="{filename}"'

    HTML(string=html_string).write_pdf(response)
    return response


# Склад временного хранения
class TSWViewSet(ModelViewSet):
    queryset = TSW.objects.all()
    serializer_class = TSWSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]


# Склад
class WarehouseViewSet(ModelViewSet):
    queryset = Warehouse.objects.all()
    serializer_class = WarehouseSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]


# Товар
class ProductViewSet(ModelViewSet):
    queryset = Product.objects.all()
    serializer_class = ProductSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]


# Рампа
class RampViewSet(ModelViewSet):
    queryset = Ramp.objects.all()
    serializer_class = RampSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]


    #http://127.0.0.1:8000/api/v1/ramps/2/occupy/ - занята рампа, /free/ - свободна
    @action(detail=True, methods=['post'], url_path='occupy')
    def occupy_ramp(self, request, pk=None):
        """
        Машина занимает рампу (статус становится 'занято').
        """
        ramp = self.get_object()
        if ramp.status_ramp == "occupied":
            return Response({"error": "Рампа уже занята."}, status=status.HTTP_400_BAD_REQUEST)

        ramp.status_ramp = "occupied"
        ramp.save()
        return Response({"message": f"Рампа {ramp.description} занята."}, status=status.HTTP_200_OK)

    @action(detail=True, methods=['post'], url_path='free')
    def free_ramp(self, request, pk=None):
        """
        Машина освобождает рампу (статус становится 'свободно').
        """
        ramp = self.get_object()
        if ramp.status_ramp == "free":
            return Response({"message": "Рампа уже свободна."}, status=status.HTTP_200_OK)

        ramp.status_ramp = "free"
        ramp.save()
        return Response({"message": f"Рампа {ramp.description} освобождена."}, status=status.HTTP_200_OK)


# Место складирования
class PlaceViewSet(ModelViewSet):
    queryset = Place.objects.all()
    serializer_class = PlaceSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]

# Тип места
class TypePlaceViewSet(ModelViewSet):
    queryset = TypePlace.objects.all()
    serializer_class = TypePlaceSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]

# Заказ услуги
class ServiceOrderViewSet(ModelViewSet):
    queryset = ServiceOrder.objects.all()
    serializer_class = ServiceOrderSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]

# Передача груза
class TransferViewSet(ModelViewSet):
    queryset = Transfer.objects.all()
    serializer_class = TransferSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]

# Заказ
class OrderViewSet(AuditLogMixin, MultiFieldFilterMixin, ModelViewSet):
    queryset = Order.objects.all().order_by('-id')  # Сортировка по id (самые ранние записи)
    serializer_class = OrderSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]

    # Указываем поля для фильтрации сделанный через миксин
    filter_fields = ['carrier_name', 'vehicle_number']  # Поля для частичного поиска
    exact_filter_fields = ['warehouse_id']  # Поля для точного поиска

    def get_serializer_class(self):
        if self.request.method in ['POST', 'PATCH']:
            return OrderCreateSerializer
        return OrderSerializer

    def perform_create(self, serializer):
        # Заполняем поле user текущим пользователем, если оно не передано
        serializer.save(user=self.request.user)


    # Использование миксина для поиска по полям
    # /orders/filter_orders/?warehouse=1
    @action(detail=False, methods=['get'])
    def filter_orders(self, request):
        return self.filter_by_fields(request)



    # запрос на поиск заказов по id и номеру авто
    # http://127.0.0.1:8000/api/v1/orders/search/?query=1

    @action(detail=False, methods=['get'], url_path='search')
    def search(self, request):
        """
        Поиск заказов по частичному совпадению с ID, vehicle_number, carrier_name, phone
        и фильтрацией по warehouse_id.
        Если параметры пустые, возвращает все записи.
        Если записи не найдены, возвращает пустой массив.
        """
        query = request.query_params.get('query', '').strip()
        warehouse_id = request.query_params.get('warehouse_id', '').strip()

        # Начинаем с полного queryset
        orders = self.queryset

        # Фильтрация по частичному совпадению ID, vehicle_number, carrier_name, phone
        if query:
            orders = orders.filter(
                Q(id__icontains=query) |
                Q(vehicle_number__iregex=query) |
                Q(carrier_name__iregex=query) |  # Поиск по имени водителя
                Q(phone__icontains=query)  # Поиск по телефону
            )

        # Фильтрация по warehouse_id
        if warehouse_id:
            orders = orders.filter(warehouse_id=warehouse_id)

        # Сортировка по id
        orders = orders.order_by('-id')

        # Пагинация
        page = self.paginate_queryset(orders)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            return self.get_paginated_response(serializer.data)

        # Если записи отсутствуют, возвращаем пустой массив
        serializer = self.get_serializer(orders, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)

    # Фильтрация по статусам status_order
    # api/orders/filtered-by-status/?status=2
    @action(detail=False, methods=['get'], url_path='filtered-by-status')
    def get_filtered_orders(self, request):
        """
        Фильтрует заказы по переданному статусу и добавляет остальные заказы сразу после.
        Используется кастомная пагинация AllOtherAPIListPagination.
        """
        filter_status = request.query_params.get('status')

        if filter_status is None:
            return Response({'error': 'Необходимо передать параметр status.'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            filter_status = str(filter_status)  # Приводим к строке, так как статусы хранятся в CharField
            filtered_orders = Order.objects.filter(status_order=filter_status)  # Фильтрованные заказы
            other_orders = Order.objects.exclude(status_order=filter_status).order_by(
                'status_order')  # Остальные заказы

            # Объединяем результаты
            orders = list(filtered_orders) + list(other_orders)

            # Применяем твою кастомную пагинацию
            paginator = AllOtherAPIListPagination()
            paginated_orders = paginator.paginate_queryset(orders, request)

            serializer = OrderSerializer(paginated_orders, many=True)
            return paginator.get_paginated_response(serializer.data)

        except Exception as e:
            return Response({'error': f'Ошибка при фильтрации: {e}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



# Парковка
class ParkingViewSet(AuditLogMixin, ModelViewSet):
    queryset = Parking.objects.all()
    serializer_class = ParkingSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]

# Парковочное место
class PlaceParkViewSet(AuditLogMixin, ModelViewSet):
    queryset = PlacePark.objects.all()
    serializer_class = PlaceParkSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]


    # Запрос на свободные парковочные места
    # /api/v1/place-parks/available/
    @action(detail=False, methods=['get'], url_path='available', url_name='available')
    def get_available_places(self, request):
        """
        Возвращает список доступных парковочных мест
        """
        available_places = PlacePark.objects.filter(is_available=True)
        serializer = self.get_serializer(available_places, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)


# Поставщик
class SupplierViewSet(AuditLogMixin, ModelViewSet):
    queryset = Supplier.objects.all()
    serializer_class = SupplierSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]


# Журнал учета
class LogBookViewSet(AuditLogMixin, MultiFieldFilterMixin, ModelViewSet):
    queryset = LogBook.objects.select_related('warehouse', 'place_park', 'user_in', 'user_out').all().order_by('-id')
    serializer_class = LogBookSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]

    # Указываем поля для фильтрации сделанный через миксин
    filter_fields = ['carrier_name', 'vehicle_number']  # Поля для частичного поиска
    exact_filter_fields = ['warehouse_id']  # Поля для точного поиска

    def get_serializer_class(self):
        if self.action == 'create':
            return LogBookCreateSerializer
        elif self.action == 'update_datetime_out':
            return LogBookUpdateSerializer
        return LogBookSerializer

    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        if serializer.is_valid():
            # Получение парковочного места
            place_park = serializer.validated_data.get('place_park')
            if place_park:
                if not place_park.is_available:
                    return Response({"error": "Парковочное место уже занято"}, status=status.HTTP_400_BAD_REQUEST)
                # Помечаем место как занятое
                place_park.is_available = False
                place_park.save()

            # Сохраняем запись и логируем
            serializer.save(user_in=request.user)
            self.perform_create(serializer)
            return Response(serializer.data, status=status.HTTP_201_CREATED)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def perform_update(self, serializer):
        instance = self.get_object()

        # Обновление состояния парковочного места, если нужно
        new_place_park = serializer.validated_data.get('place_park')
        if new_place_park and new_place_park != instance.place_park:
            if not new_place_park.is_available:
                raise ValidationError("Новое парковочное место уже занято")
            if instance.place_park:
                instance.place_park.is_available = True
                instance.place_park.save()
            new_place_park.is_available = False
            new_place_park.save()

        # Сохраняем и логируем
        super().perform_update(serializer)

    def perform_destroy(self, instance):
        # Если требуется освободить парковочное место
        if instance.place_park:
            instance.place_park.is_available = True
            instance.place_park.save()

        # Удаляем объект и логируем
        super().perform_destroy(instance)



    # Запрос через миксин на поиск по полям
    # api/v1/log-books/filter_orders/?warehouse_id=2&vehicle_number=
    @action(detail=False, methods=['get'])
    def filter_orders(self, request):
        """
        Фильтрация записей LogBook по указанным параметрам.
        """
        # Проверяем, переданы ли параметры фильтрации
        if not request.query_params:
            return Response([], status=200)  # Пустой массив

        # Проверяем, что значения всех параметров не пустые
        for key, value in request.query_params.items():
            if value == "":
                return Response([], status=200)  # Пустой массив

        # Используем миксин для получения отфильтрованного queryset
        queryset = self.filter_by_fields(request)

        if not queryset.exists():
            return Response([], status=200)  # Пустой массив

        # Пагинация и сериализация результата
        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            return self.get_paginated_response(serializer.data)

        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)


    # PATCH /api/v1/log-books/1/update-out
    # {
    #     "datetime_out": "2024-12-03T18:00:00",
    #     "removed_control": "yes",
    #     "note": "Машина выехала"
    # }

    @action(detail=True, methods=['patch'], url_path='update-out')
    def update_datetime_out(self, request, pk=None):
        """
        Обновляет `datetime_out`, устанавливает `user_out` и освобождает место.
        """
        try:
            log_entry = self.get_object()
        except LogBook.DoesNotExist:
            return Response({"error": "Запись не найдена"}, status=status.HTTP_404_NOT_FOUND)

        serializer = self.get_serializer(log_entry, data=request.data, partial=True)
        if serializer.is_valid():
            # Освобождаем место, если связано с записью
            if log_entry.place_park:
                log_entry.place_park.is_available = True
                log_entry.place_park.save()

            # Устанавливаем `user_out` как текущего пользователя
            serializer.save(user_out=request.user)

            return Response(LogBookSerializer(log_entry).data, status=status.HTTP_200_OK)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


    # запрос на поиск заказов по id и номеру авто
    # http://127.0.0.1:8000/api/v1/log-books/search/?query=1
    @action(detail=False, methods=['get'], url_path='search')
    def search(self, request):
        """
        Поиск заказов по частичному совпадению с ID, vehicle_number и фильтрацией по warehouse_id.
        Если параметры пустые, возвращает все записи.
        Если записи не найдены, возвращает пустой массив.
        """
        query = request.query_params.get('query', '').strip()
        warehouse_id = request.query_params.get('warehouse_id', '').strip()

        # Начинаем с полного queryset
        logbooks = self.queryset

        # Фильтрация по частичному совпадению ID или vehicle_number
        if query:
            logbooks = logbooks.filter(
                Q(id__icontains=query) | Q(vehicle_number__icontains=query)
            )

        # Фильтрация по warehouse_id
        if warehouse_id:
            logbooks = logbooks.filter(warehouse_id=warehouse_id)

        # Сортировка по id
        logbooks = logbooks.order_by('-id')

        # Пагинация
        page = self.paginate_queryset(logbooks)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            return self.get_paginated_response(serializer.data)

        # Если записи отсутствуют, возвращаем пустой массив
        serializer = self.get_serializer(logbooks, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)


# Уведомление
class NoticeViewSet(AuditLogMixin, ModelViewSet):
    queryset = Notice.objects.all()
    serializer_class = NoticeSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]

    def perform_create(self, serializer):
        # Передаем текущего пользователя в сериализатор
        serializer.save(user=self.request.user)

    def get_serializer_class(self):
        """
        Возвращаем другой сериализатор для PATCH запросов с исключением поля Order
        """
        if self.action == 'partial_update':  # PATCH запрос
            return NoticePatchSerializer
        return NoticeSerializer


    # Запрос на создание черновика с постоянными данными в таблице уведомление
    #/ api / v1 / notices / create - empty /
    # POST http://127.0.0.1:8000/api/v1/notices/create-empty/
    # {
    #   "order": 1
    # }
    @action(detail=False, methods=['post'], url_path='create-empty')
    def create_empty_notice(self, request):
        """
        Создает черновик уведомления с фиксированными данными, включая данные из Transport.
        """
        user = request.user
        order_id = request.data.get('order')

        if not order_id:
            return Response({'error': 'Необходимо указать заказ.'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            order = Order.objects.select_related('warehouse', 'logbook').get(id=order_id)
        except Order.DoesNotExist:
            return Response({'error': 'Заказ не найден.'}, status=status.HTTP_404_NOT_FOUND)

        # Получаем данные из Order
        vehicle_number = order.vehicle_number  # Номер машины из Order

        # Получаем данные из LogBook, включая trailer_number
        logbook = order.logbook
        trailer_number = logbook.trailer_number if logbook else None

        # Получаем данные из Transport по номеру машины
        transport = Transport.objects.filter(ts=vehicle_number).first()
        transport_data = {
            "ts": transport.ts if transport else None,
            "type_ts": transport.type_ts if transport else None,
            "country": transport.country if transport else None
        }

        # Если есть trailer_number, проверяем его в Transport
        trailer_data = {}
        if trailer_number:
            trailer_transport = Transport.objects.filter(ts=trailer_number).first()
            if trailer_transport:
                trailer_data = {
                    "trailer_ts": trailer_transport.ts,
                    "trailer_type_ts": trailer_transport.type_ts,
                    "trailer_country": trailer_transport.country
                }

        # Генерация GUID
        guid = str(uuid.uuid4()).upper()

        # Получаем номер склада из связанной таблицы Warehouse
        svh_number = order.warehouse.svh_number if order.warehouse else None

        # Создаем уведомление
        notice = Notice.objects.create(
            user=user,
            order=order,
            stz="0",
            zhurnal="9",
            year=str(datetime.now().year)[-1],
            guid=guid,
            date_in=logbook.datetime_in.date() if logbook else None,  # Устанавливаем дату из LogBook
            time_in=logbook.datetime_in.time() if logbook else None,  # Устанавливаем время из LogBook
        )

        # Устанавливаем дополнительные значения
        notice.number_out = str(notice.id)
        notice.notification = str(notice.id).zfill(6)  # Должно быть 6 символов
        notice.number_notification = f"СВ-1601/0000304/{notice.year}{notice.zhurnal}{notice.notification}"
        notice.save()
        # notice.number_out = str(notice.id)
        # notice.notification = f"{'0' * (5 - len(str(notice.id)))}{notice.id}"
        # notice.number_notification = f"СВ-1601/0000304/{notice.year}{notice.zhurnal}{notice.stz}{str(notice.id).zfill(5)}"
        # notice.save()

        # Возвращаем заполненное уведомление
        serializer = self.get_serializer(notice)
        response_data = serializer.data
        response_data.update(transport_data)
        response_data.update(trailer_data)  # Добавляем данные прицепа, если они есть
        response_data['svh_number'] = svh_number

        return Response(response_data, status=status.HTTP_201_CREATED)




    # Поиск уведолмения по id заказа
    # /api/v1/notices/by-order/1/
    # @action(detail=False, methods=['get'], url_path='by-order/(?P<order_id>[^/.]+)', url_name='by-order')
    # def get_notices_by_order_id(self, request, order_id=None):
    #     """
    #     Получает уведомления по ID заказа вместе с документами, получателями, транспортом и складом.
    #     """
    #     try:
    #         # Получаем уведомления по ID заказа
    #         notices = Notice.objects.filter(order_id=order_id).select_related('order__warehouse').prefetch_related(
    #             'docs', 'recipient')
    #         if not notices.exists():
    #             return Response(
    #                 {"error": f"Уведомления для заказа с ID {order_id} не найдены."},
    #                 status=status.HTTP_404_NOT_FOUND,
    #             )
    #
    #         serializer = NoticeSerializer(notices, many=True)
    #         return Response(serializer.data, status=status.HTTP_200_OK)
    #     except Exception as e:
    #         return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)
    @action(detail=False, methods=['get'], url_path='by-order/(?P<order_id>[^/.]+)', url_name='by-order')
    def get_notices_by_order_id(self, request, order_id=None):
        """
        Получает уведомления по ID заказа вместе с документами, получателями, транспортом, складом и статусом EClient.
        """
        try:
            # Получаем уведомления по ID заказа
            notices = Notice.objects.filter(order_id=order_id).select_related('order__warehouse').prefetch_related(
                'docs', 'recipient'
            )

            if not notices.exists():
                return Response(
                    {"error": f"Уведомления для заказа с ID {order_id} не найдены."},
                    status=status.HTTP_404_NOT_FOUND,
                )

            # Сериализуем данные (теперь status подтянется автоматически)
            serializer = NoticeSerializer(notices, many=True)
            return Response(serializer.data, status=status.HTTP_200_OK)

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)




# Документ
class DocViewSet(ModelViewSet):
    queryset = Doc.objects.all()
    serializer_class = DocSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]

    def create(self, request, *args, **kwargs):
        """
        Обрабатываем POST-запрос для создания документа.
        """
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)

        # Сохраняем документ и возвращаем notice_id
        doc = serializer.save()
        response_data = serializer.data
        response_data['notice_id'] = doc.notice.id  # Добавляем поле notice_id
        return Response(response_data, status=status.HTTP_201_CREATED)



class DocumentFileViewSet(AuditLogMixin, MultiFieldFilterMixin, ModelViewSet):
    queryset = DocumentFile.objects.all()
    serializer_class = DocumentFileSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser]

    # Настройка фильтров api/v1/document-files/?order=1
    filter_fields = []  # Поля для частичного поиска (если потребуется)
    exact_filter_fields = ['order']  # Поля для точного поиска

    def get_queryset(self):
        # Получаем исходный queryset
        base_queryset = super().get_queryset()
        # Применяем фильтрацию из миксина
        return self.filter_by_fields(self.request, queryset=base_queryset)



   #  Скачивание целиком папки в zip по ID заказу и отдельно по файлу
   #  http://127.0.0.1:8000/api/v1/document-files/download/2/?file_id=1
    @action(detail=False, methods=['get'], url_path=r'download/(?P<order_id>\d+)')
    def download_files(self, request, order_id=None):
        # Получаем все файлы для указанного заказа
        files = DocumentFile.objects.filter(order_id=order_id)

        if not files:
            raise NotFound("Файлы для данного заказа не найдены.")

        # Если передан параметр file_id, скачиваем конкретный файл
        file_id = request.query_params.get('file_id')
        if file_id:
            try:
                document_file = files.get(id=file_id)
                return self.download_file(document_file)
            except DocumentFile.DoesNotExist:
                raise NotFound("Файл не найден.")

        # Если file_id не передан, создаем и отправляем архив с файлами для данного заказа
        return self.download_all_files_as_zip(files, order_id)

    def download_file(self, document_file):
        # Проверяем, существует ли файл
        file_path = document_file.file.path
        if not os.path.exists(file_path):
            raise NotFound("Файл не найден на сервере.")

        # Подготавливаем имя файла
        file_name = document_file.get_file_name()

        try:
            # Открываем файл для скачивания
            with open(file_path, 'rb') as file:
                response = HttpResponse(file.read(), content_type='application/octet-stream')
                response['Content-Disposition'] = f'attachment; filename="{file_name}"'
                return response
        except FileNotFoundError:
            raise NotFound(f"Файл {file_name} не найден на сервере.")

    def download_all_files_as_zip(self, files, order_id):
        # Создаем временный архив для скачивания
        zip_filename = f"{order_id}.zip"  # Имя архива = order_id.zip
        zip_file_path = os.path.join(settings.MEDIA_ROOT, zip_filename)

        try:
            with zipfile.ZipFile(zip_file_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for document_file in files:
                    file_path = document_file.file.path
                    arcname = os.path.join(f"order_{order_id}",
                                           document_file.get_file_name())  # Папка с именем order_id
                    zipf.write(file_path, arcname=arcname)
        except Exception as e:
            raise NotFound(f"Ошибка при создании архива: {e}")

        # Отправляем архив для скачивания
        with open(zip_file_path, 'rb') as zip_file:
            response = HttpResponse(zip_file.read(), content_type='application/zip')
            response['Content-Disposition'] = f'attachment; filename={zip_filename}'

        # Удаляем временный архив после отправки
        os.remove(zip_file_path)

        return response



# Информация о перемещении товара по местам
class LogPlaceViewSet(ModelViewSet):
    queryset = LogPlace.objects.select_related('product', 'place_from', 'place_to', 'user').all()
    serializer_class = LogPlaceSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]



# Получатель
class RecipientViewSet(ModelViewSet):
    queryset = Recipient.objects.all()
    serializer_class = RecipientSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]


# Транспорт
class TransportViewSet(AuditLogMixin, ModelViewSet):
    queryset = Transport.objects.all().order_by('-id')
    serializer_class = TransportSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]

    # Поиск по t/s
    # http://127.0.0.1:8000/transports/search_by_ts/?q=4
    @action(detail=False, methods=['get'])
    def search_by_ts(self, request):
        search_query = request.query_params.get('q', '').strip()
        if not search_query:
            return Response({"error": "Неверный параметр ввода"}, status=400)

        # Выполняем фильтрацию
        transports = self.queryset.filter(Q(ts__iregex=search_query)).order_by('-id')[:6]

        # Проверяем результат
        if not transports.exists():
            return Response({"message": "false"}, status=404)

        # Сериализуем результат
        serializer = self.get_serializer(transports, many=True)
        return Response(serializer.data)

    # Поиск по t/s и по carrier_name
    # http://127.0.0.1:8000/transports/search_by_ts_and_carrier_name/?q=4
    @action(detail=False, methods=['get'])
    def search_by_ts_and_carrier_name(self, request):
        search_query = request.query_params.get('q', '').strip()

        # Если параметр q пуст, возвращаем все записи
        if not search_query:
            transports = self.queryset
        else:
            # Фильтрация по ts и carrier_name
            transports = self.queryset.filter(
                Q(ts__iregex=search_query) | Q(carrier_name__iregex=search_query)
            )

        # Проверяем результат
        if not transports.exists():
            return Response({"message": "false"}, status=404)

        # Применяем пагинацию через AllOtherAPIListPagination
        page = self.paginate_queryset(transports)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            return self.get_paginated_response(serializer.data)

        # Сериализуем результат без пагинации
        serializer = self.get_serializer(transports, many=True)
        return Response(serializer.data)







# Транспорт-Уведомление
class TransportNoticeViewSet(ModelViewSet):
    queryset = TransportNotice.objects.all()
    serializer_class = TransportNoticeSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]




# Запрос на справочники
# http://127.0.0.1:8000/api/v1/directories/purpose_types/
# Справочники находятся в directory.py
class DirectoryView(APIView):
    """
    API для предоставления справочников в формате массива объектов.
    """
    def get(self, request, *args, **kwargs):
        directory_name = kwargs.get('directory_name')
        directories = {
            'http_answer': HTTP_ANSWER,
            'languages': LANG,
            'status_reg': STATUS_REG,
            'ln_type': LN_TYPE,
            'countries': COUNTRY,
            'currencies': CURRENCY_CHOICES,
            'purpose_types': PURPOSE_TYPES,
            'type_service': TYPE_SERVICE,
            'type_ts': TYPE_TS,
            'type_ts_choises': TRANSPORT_TYPE_CHOICES,
            'direction_types': DIRECTION_TYPES,
            'status_order': STATUS_ORDER,


        }

        if directory_name in directories:
            data = directories[directory_name]

            # Преобразуем список пар в массив объектов
            result = [{"code": item[0], "name": item[1]} for item in data]

            return Response(result, status=status.HTTP_200_OK)
        else:
            return Response(
                {"error": f"Справочник '{directory_name}' не найден."},
                status=status.HTTP_404_NOT_FOUND
            )



class UserActionLogViewSet(ReadOnlyModelViewSet):
    """
    Представление для просмотра логов действий пользователей с возможностью фильтрации.
    """
    queryset = UserActionLog.objects.all().order_by('-datetime')
    serializer_class = UserActionLogSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAdminUser]


    # Фильтрация по дате и действиям
    # http://127.0.0.1:8000/api/v1/user-actions/?start_date=2024-02-01&end_date=2024-02-05
    # http://178.168.146.114:8770/api/v1/user-actions/?action=create
    # action=update
    # action=delete
    def get_queryset(self):
        """
        Фильтрует логи по параметрам action и дате.
        """
        queryset = super().get_queryset()
        action = self.request.query_params.get('action')
        start_date = self.request.query_params.get('start_date')
        end_date = self.request.query_params.get('end_date')

        # Фильтрация по action
        if action:
            queryset = queryset.filter(action=action)

        # Фильтрация по start_date
        if start_date:
            try:
                # Преобразуем start_date в UTC-aware datetime
                start_date_obj = make_aware(datetime.strptime(start_date, "%Y-%m-%d"), utc)
                queryset = queryset.filter(datetime__gte=start_date_obj)
            except ValueError:
                pass

        # Фильтрация по end_date
        if end_date:
            try:
                # Преобразуем end_date в конец дня и делаем его UTC-aware
                end_date_obj = make_aware(datetime.strptime(end_date, "%Y-%m-%d") + timedelta(days=1), utc)
                queryset = queryset.filter(datetime__lt=end_date_obj)
            except ValueError:
                pass

        return queryset


class EClientViewSet(ModelViewSet):
    """
    ViewSet для таблицы eclient из второй базы данных.
    """
    queryset = EClient.objects.using('second_db').all().order_by("-ideclient")
    serializer_class = EClientSerializer
    pagination_class = AllOtherAPIListPagination
    permission_classes = [IsAuthenticated]

    # Поиск по номеру авто и ID регистрации по частичному поиску + фильтрация по дате по полю lc
    # /api/v1/eclients/search/?query=2
    # /api/v1/eclients/search/?start_date=01.02.2024&end_date=10.02.2024
    @action(detail=False, methods=['get'], url_path='search')
    def search_by_number(self, request):
        """
        Поиск по `ideclient`, `regN_TS` и `lc`
        """
        search_query = request.query_params.get('query', None)
        start_date = request.query_params.get('start_date', None)
        end_date = request.query_params.get('end_date', None)

        if not start_date or not end_date:
            return Response({"error": "Введите start_date и end_date в формате YYYY-MM-DD"}, status=400)

        queryset = self.queryset

        # Фильтрация по query (если указан)
        if search_query and search_query.strip():
            queryset = queryset.filter(
                Q(ideclient__icontains=search_query) |
                Q(regN_TS__icontains=search_query)
            )

        # Фильтрация по дате lc (учитываем, что в БД она как datetime!)
        try:
            if start_date:
                start_date_parsed = make_aware(datetime.strptime(start_date, "%Y-%m-%d"))
                queryset = queryset.filter(lc__gte=start_date_parsed)

            if end_date:
                end_date_parsed = make_aware(datetime.strptime(end_date, "%Y-%m-%d")) + timedelta(
                    days=1)  # +1 день для включения границы
                queryset = queryset.filter(lc__lt=end_date_parsed)

        except ValueError:
            return Response({"error": "Неверный формат дат. Используйте YYYY-MM-DD"}, status=400)

        # Проверяем, есть ли записи
        if not queryset.exists():
            return Response({"message": "Записи за указанный период отсутствуют"}, status=404)

        # Пагинация
        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            return self.get_paginated_response(serializer.data)

        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data, status=200)

    # Создание отчета по одному полю regID в DOCX
    # GET /api/v1/eclients/generate-report/?start_date=01.02.2025&end_date=20.02.2025
    @action(detail=False, methods=['get'], url_path='generate-report')
    def generate_docx_report(self, request):
        """
        Генерация отчета в формате DOCX (без подсчета записей и без пустых значений).
        """
        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')

        if not start_date or not end_date:
            return Response({"error": "Укажите start_date и end_date в формате YYYY-MM-DD"}, status=400)

        try:
            start_date_parsed = datetime.strptime(start_date, "%Y-%m-%d")
            end_date_parsed = datetime.strptime(end_date, "%Y-%m-%d")
        except ValueError:
            return Response({"error": "Неверный формат дат. Используйте YYYY-MM-DD"}, status=400)

        # Преобразуем дату в формат DD.MM.YYYY для отображения в документе и названии файла
        start_date_display = start_date_parsed.strftime("%d.%m.%Y")
        end_date_display = end_date_parsed.strftime("%d.%m.%Y")

        # Фильтрация записей
        records = self.queryset.filter(
            lc__gte=start_date_parsed,
            lc__lte=end_date_parsed
        ).exclude(numbEPI__isnull=True).exclude(numbEPI__exact="")

        if not records.exists():
            return Response({"error": "Записи за указанный период отсутствуют"}, status=404)

        # Сбор данных без пустых значений
        numb_epi_list = [numb_epi for numb_epi in records.values_list('numbEPI', flat=True)]

        warehouse = Warehouse.objects.first()
        svh_number = warehouse.svh_number if warehouse else "СВ-1601/0000304"

        # Создание документа
        doc = Document()

        # Заголовок отчета
        title = doc.add_paragraph(
            "Отчет об уведомлениях о размещении товаров в зоне таможенного контроля, снятых с учета")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].bold = True
        title.runs[0].font.size = Pt(14)

        # Период в документе (формат DD.MM.YYYY)
        subtitle = doc.add_paragraph(f"за период с {start_date_display} по {end_date_display}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(12)

        # Номер СВХ
        svh_paragraph = doc.add_paragraph(svh_number)
        svh_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        svh_paragraph.runs[0].font.size = Pt(12)

        # Описание СВХ
        svh_desc_paragraph = doc.add_paragraph(
            "(номер включения в соответствующий реестр владельца склада временного хранения "
            "или свободного склада или регистрационный номер зоны таможенного контроля по реестру зон "
            "таможенного контроля, формируемому таможенным органом, в регионе деятельности которого "
            "создана зона таможенного контроля)"
        )
        svh_desc_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        svh_desc_paragraph.runs[0].italic = True
        svh_desc_paragraph.runs[0].font.size = Pt(10)

        # Создание таблицы
        table = doc.add_table(rows=2, cols=1)
        table.style = "Table Grid"

        # Заголовок таблицы
        header_row = table.rows[0].cells[0]
        header_paragraph = header_row.paragraphs[0]
        header_paragraph.add_run(
            "Регистрационный номер уведомления о размещении товаров в зоне таможенного контроля").bold = True
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Подзаголовок строки
        subheader_row = table.rows[1].cells[0]
        subheader_paragraph = subheader_row.paragraphs[0]
        subheader_paragraph.add_run("1").bold = True
        subheader_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Добавление данных
        for numb_epi in numb_epi_list:
            row = table.add_row().cells[0]
            row_paragraph = row.paragraphs[0]
            row_paragraph.add_run(f"{numb_epi}")
            row_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Сохранение документа в память
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Русское имя файла (формат DD.MM.YYYY)
        filename = f"Отчет по уведомлениям за период {start_date_display} - {end_date_display}.docx"
        encoded_filename = urllib.parse.quote(filename)

        # Возврат документа с русским именем файла
        return HttpResponse(
            buffer,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f'attachment; filename*=UTF-8\'\'{encoded_filename}'
            }
        )

class CustomTokenObtainPairView(TokenObtainPairView):
    serializer_class = CustomTokenObtainPairSerializer







# URL API
BTS_URL = 'http://web.declarant.by/api/gsec-proxy/ServiceISZL/ecd/v1'

# Твои данные авторизации (замени на актуальные)
AUTH_TOKEN = "93dc1f95-7902-3d17-baa3-fa6387d99076"
USER_ID = "0D1D389D-EBA8-41C2-94A9-4A66817E21A8"


def get_headers():
    """Создает заголовки для запроса"""
    return {
        'Authorization': f'Bearer {AUTH_TOKEN}',
        'UserId': USER_ID,
        'isBTS': 'True',
        'Content-Type': 'application/xml',
        'Accept-Charset': 'utf-8'
    }





@csrf_exempt
def get_registration_info(request, guid):
    """Запрос информации по регистрации"""
    url = f"{BTS_URL}/requests?file_guid={guid}&limit=10"

    try:
        response = requests.get(url, headers=get_headers())

        if response.status_code == 200:
            return JsonResponse(response.json(), status=200)

        return JsonResponse({"error": response.json()}, status=response.status_code)

    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)



@csrf_exempt
def get_file_list(request, guid):
    """Запрос списка файлов по guid (если id_reg нет, сначала получаем его), возвращаем JSON"""

    try:
        with transaction.atomic():
            # Проверяем, есть ли запись с таким guidXML в eclient
            try:
                eclient = EClient.objects.get(guidXML=guid)
                id_reg = eclient.regID  # Берем id_reg из regID
            except EClient.DoesNotExist:
                id_reg = None

            # Если id_reg нет, сначала получаем его через API
            if not id_reg:
                reg_url = f"{BTS_URL}/requests?file_guid={guid}&limit=10"
                reg_response = requests.get(reg_url, headers=get_headers())

                if reg_response.status_code == 200:
                    reg_data = reg_response.json()

                    if "requests" in reg_data and len(reg_data["requests"]) > 0:
                        id_reg = reg_data["requests"][0]["id"]  # Получаем ID из API

                        # Записываем ID в БД
                        eclient, created = EClient.objects.get_or_create(guidXML=guid)
                        eclient.regID = id_reg
                        eclient.save()
                    else:
                        return JsonResponse({"error": reg_data}, status=400)
                else:
                    return JsonResponse({"error": reg_response.json()}, status=reg_response.status_code)

            # Делаем запрос на файлы по id_reg
            file_url = f"{BTS_URL}/files/{id_reg}"
            file_response = requests.get(file_url, headers=get_headers())

            if file_response.status_code == 200:
                return JsonResponse(file_response.json(), status=200)  # Возвращаем JSON с файлами

            return JsonResponse({"error": file_response.json()}, status=file_response.status_code)

    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)
# @csrf_exempt
# def get_file_list(request, guid):
#     """Запрос списка файлов по guid (если id_reg нет, сначала получаем его)"""
#
#     try:
#         with transaction.atomic():
#             # Ищем запись с таким guidXML
#             try:
#                 client = EClient.objects.get(guidXML=guid)
#                 id_reg = client.regID  # id_reg берем из regID
#             except EClient.DoesNotExist:
#                 id_reg = None
#
#             # Если id_reg не найден, делаем первый запрос, чтобы его получить
#             if not id_reg:
#                 reg_url = f"{BTS_URL}/requests?file_guid={guid}&limit=10"
#                 reg_response = requests.get(reg_url, headers=get_headers())
#
#                 print("СТАТУС ПОЛУЧЕНИЯ ID:", reg_response.status_code)
#                 print("ОТВЕТ ПОЛУЧЕНИЯ ID:", reg_response.text[:500])
#
#                 if reg_response.status_code == 200:
#                     reg_data = reg_response.json()
#
#                     if "requests" in reg_data and len(reg_data["requests"]) > 0:
#                         id_reg = reg_data["requests"][0]["id"]  # Получаем ID из API
#
#                         # Записываем ID в БД
#                         client, created = EClient.objects.get_or_create(guidXML=guid)
#                         client.regID = id_reg
#                         client.save()
#                     else:
#                         return JsonResponse({"error": "Не удалось получить ID по GUID"}, status=400)
#                 else:
#                     return JsonResponse(
#                         {'error': f'Ошибка при получении ID: {reg_response.status_code}, {reg_response.text}'},
#                         status=reg_response.status_code)
#
#             # Делаем запрос на файлы по id_reg
#             file_url = f"{BTS_URL}/files/{id_reg}"
#             file_response = requests.get(file_url, headers=get_headers())
#
#             print("СТАТУС ПОЛУЧЕНИЯ ФАЙЛОВ:", file_response.status_code)
#             print("ОТВЕТ ФАЙЛОВ:", file_response.text[:500])
#
#             if file_response.status_code == 200:
#                 return JsonResponse(file_response.json(), status=200)
#
#             return JsonResponse(
#                 {'error': f'Ошибка при получении файлов: {file_response.status_code}, {file_response.text}'},
#                 status=file_response.status_code)
#
#     except Exception as e:
#         return JsonResponse({'error': f'Exception: {str(e)}'}, status=500)
# @csrf_exempt
# def get_file_list(request, id_reg):
#     """Запрос списка файлов по id_reg"""
#     url = f"{BTS_URL}/files/{id_reg}"
#
#     try:
#         response = requests.get(url, headers=get_headers())
#
#         # Логируем ответ API
#         print("СТАТУС:", response.status_code)
#         print("ОТВЕТ:", response.text)
#
#         if response.status_code == 200:
#             return JsonResponse(response.json(), status=200)
#         return JsonResponse({'error': f'Ошибка: {response.status_code}, {response.text}'}, status=response.status_code)
#
#     except Exception as e:
#         return JsonResponse({'error': f'Exception: {str(e)}'}, status=500)


@csrf_exempt
def get_file(request, file_id):
    """Запрос файла по file_id и возврат XML в браузере"""
    url = f"{BTS_URL}/file/{file_id}"

    try:
        response = requests.get(url, headers=get_headers())

        if response.status_code == 200:
            content_type = response.headers.get("Content-Type", "")

            # Если это XML или бинарный файл, возвращаем его сразу
            if "application/xml" in content_type or "application/octet-stream" in content_type:
                return HttpResponse(response.content, content_type="application/xml")

            # Если это JSON, просто отдаем его
            if "application/json" in content_type:
                return JsonResponse(response.json(), status=200)

        # Обрабатываем ошибку API
        return JsonResponse({
            "error": "Ошибка при получении файла",
            "status": response.status_code,
            "details": response.json() if "application/json" in response.headers.get("Content-Type", "") else response.text
        }, status=response.status_code)

    except Exception as e:
        return JsonResponse({
            "error": str(e),
            "details": "Ошибка на стороне сервера"
        }, status=500)